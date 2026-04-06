"""
Servicio para extraer texto de diferentes formatos de archivo.
Soporta: PDF, DOC, DOCX, TXT
"""
import io
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document


class FileParserError(Exception):
    """Excepción personalizada para errores de parseo."""
    pass


class FileParser:
    """Parser de archivos a texto."""
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MAX_CHARS = 50000  # Máximo caracteres para TTS
    
    ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}
    
    @classmethod
    def validate_file(cls, filename: str, content: bytes) -> None:
        """Valida que el archivo sea permitido."""
        import os
        
        ext = os.path.splitext(filename.lower())[1]
        
        if ext not in cls.ALLOWED_EXTENSIONS:
            raise FileParserError(f"Formato no soportado: {ext}. Use: PDF, DOC, DOCX, TXT")
        
        if len(content) > cls.MAX_FILE_SIZE:
            raise FileParserError(f"Archivo muy grande. Máximo: {cls.MAX_FILE_SIZE // 1024 // 1024}MB")
    
    @classmethod
    def extract_text(cls, filename: str, content: bytes) -> str:
        """Extrae texto del archivo según su formato."""
        import os
        
        cls.validate_file(filename, content)
        
        ext = os.path.splitext(filename.lower())[1]
        
        try:
            if ext == '.pdf':
                text = cls._extract_pdf(content)
            elif ext == '.docx':
                text = cls._extract_docx(content)
            elif ext == '.doc':
                text = cls._extract_doc(content)
            elif ext == '.txt':
                text = cls._extract_txt(content)
            else:
                raise FileParserError(f"Formato no implementado: {ext}")
            
            # Limpiar y limitar texto
            text = cls._clean_text(text)
            
            if len(text) > cls.MAX_CHARS:
                text = text[:cls.MAX_CHARS]
                text += "\n\n[Texto truncado por exceder límite de 50,000 caracteres]"
            
            if len(text.strip()) < 10:
                raise FileParserError("El archivo no contiene texto suficiente (mínimo 10 caracteres)")
            
            return text
            
        except FileParserError:
            raise
        except Exception as e:
            raise FileParserError(f"Error al procesar archivo: {str(e)}")
    
    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """Extrae texto de PDF."""
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            raise FileParserError(f"Error leyendo PDF: {str(e)}")
    
    @staticmethod
    def _extract_docx(content: bytes) -> str:
        """Extrae texto de DOCX."""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # También extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            raise FileParserError(f"Error leyendo DOCX: {str(e)}")
    
    @staticmethod
    def _extract_doc(content: bytes) -> str:
        """Extrae texto de DOC (formato antiguo)."""
        # Los archivos .doc son más difíciles. Intentamos con textract o damos mensaje
        try:
            # Intentar leer como texto plano primero
            text = content.decode('utf-8', errors='ignore')
            # Limpiar caracteres no imprimibles comunes en DOC
            text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
            
            if len(text.strip()) > 100:
                return text
            
            raise FileParserError(
                "Archivos .doc antiguos requieren conversión. "
                "Por favor, guárdelo como .docx o PDF y vuelva a subirlo."
            )
        except Exception as e:
            raise FileParserError(f"Error leyendo DOC: {str(e)}")
    
    @staticmethod
    def _extract_txt(content: bytes) -> str:
        """Extrae texto de archivo TXT."""
        # Intentar diferentes codificaciones
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Si ninguna funciona, usar latin-1 con reemplazo
        return content.decode('latin-1', errors='replace')
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Limpia el texto extraído."""
        # Eliminar caracteres nulos y de control excepto saltos de línea
        text = ''.join(char for char in text if char == '\n' or char == '\r' or char == '\t' or (char.isprintable() and ord(char) >= 32))
        
        # Normalizar saltos de línea
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Eliminar líneas vacías múltiples (más de 2 seguidas)
        lines = text.split('\n')
        cleaned_lines = []
        empty_count = 0
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                empty_count += 1
                if empty_count <= 2:
                    cleaned_lines.append('')
            else:
                empty_count = 0
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()